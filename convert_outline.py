#!/usr/bin/env python3
"""Convert thesis-outline.md to a clean DOCX with auto-numbered headings."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
import html


def setup_multilevel_numbering(doc):
    """Set up multi-level list numbering linked to Heading 1/2/3 styles.

    This creates a numbering definition so that:
    - Heading 1 → "1", "2", "3", ...
    - Heading 2 → "1.1", "1.2", ...
    - Heading 3 → "1.1.1", "1.1.2", ...

    When sections are reorganized, Word auto-updates the numbers.
    """
    # Get or create the numbering part
    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part.numbering_definitions._numbering

    # Find the max abstractNumId and numId already in use
    max_abstract = -1
    max_num = -1
    for ab in numbering_elem.findall(qn('w:abstractNum')):
        aid = int(ab.get(qn('w:abstractNumId')))
        if aid > max_abstract:
            max_abstract = aid
    for n in numbering_elem.findall(qn('w:num')):
        nid = int(n.get(qn('w:numId')))
        if nid > max_num:
            max_num = nid

    abstract_num_id = max_abstract + 1
    num_id = max_num + 1

    # Create abstract numbering definition with 3 levels
    abstract_num_xml = f'''
    <w:abstractNum w:abstractNumId="{abstract_num_id}" {nsdecls('w')}>
        <w:multiLevelType w:val="multilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:pStyle w:val="Heading1"/>
            <w:lvlText w:val="%1"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="432" w:hanging="432"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:pStyle w:val="Heading2"/>
            <w:lvlText w:val="%1.%2"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="576" w:hanging="576"/>
            </w:pPr>
        </w:lvl>
        <w:lvl w:ilvl="2">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:pStyle w:val="Heading3"/>
            <w:lvlText w:val="%1.%2.%3"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="720" w:hanging="720"/>
            </w:pPr>
        </w:lvl>
    </w:abstractNum>
    '''

    abstract_elem = parse_xml(abstract_num_xml)
    numbering_elem.append(abstract_elem)

    # Create num element referencing the abstract
    num_xml = f'''
    <w:num w:numId="{num_id}" {nsdecls('w')}>
        <w:abstractNumId w:val="{abstract_num_id}"/>
    </w:num>
    '''
    num_elem = parse_xml(num_xml)
    numbering_elem.append(num_elem)

    return num_id


def link_heading_style_to_numbering(doc, style_name, num_id, ilvl):
    """Link a heading style to our numbering definition at the given level."""
    style = doc.styles[style_name]
    pPr = style.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        style.element.insert(0, pPr)

    # Remove existing numPr if any
    for old in pPr.findall(qn('w:numPr')):
        pPr.remove(old)

    numPr = parse_xml(f'''
        <w:numPr {nsdecls('w')}>
            <w:ilvl w:val="{ilvl}"/>
            <w:numId w:val="{num_id}"/>
        </w:numPr>
    ''')
    pPr.append(numPr)


def strip_section_number(text):
    """Remove leading section numbers like '2.1', '2.1.1', '4.6.3' from heading text."""
    return re.sub(r'^\d+(\.\d+)*\s+', '', text)


def strip_chapter_prefix(text):
    """Remove 'Chapter N:' prefix from chapter headings."""
    return re.sub(r'^Chapter\s+\d+:\s*', '', text)


def parse_inline(text):
    """Parse inline markdown and HTML into a list of (text, style_dict) tuples.

    Handles: **bold**, *italic*, `code`, <b>bold</b>,
    <span style="color:rgb(...)">text</span>, [TAGS]
    """
    segments = []

    # First, handle HTML spans with color
    # We'll do a manual parse to handle nested elements

    # Simple approach: process the string character by character isn't practical.
    # Use regex-based approach instead.

    def process_segment(s, base_style=None):
        """Recursively process a text segment into styled parts."""
        if base_style is None:
            base_style = {}

        result = []

        # Handle <span style="color:rgb(...)">...</span>
        span_pattern = r'<span\s+style="color:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)">(.*?)</span>'
        match = re.search(span_pattern, s, re.DOTALL)
        if match:
            before = s[:match.start()]
            r, g, b = int(match.group(1)), int(match.group(2)), int(match.group(3))
            inner = match.group(4)
            after = s[match.end():]

            if before:
                result.extend(process_segment(before, base_style))
            color_style = {**base_style, 'color': (r, g, b)}
            result.extend(process_segment(inner, color_style))
            if after:
                result.extend(process_segment(after, base_style))
            return result

        # Handle <b>...</b>
        b_pattern = r'<b>(.*?)</b>'
        match = re.search(b_pattern, s, re.DOTALL)
        if match:
            before = s[:match.start()]
            inner = match.group(1)
            after = s[match.end():]

            if before:
                result.extend(process_segment(before, base_style))
            bold_style = {**base_style, 'bold': True}
            result.extend(process_segment(inner, bold_style))
            if after:
                result.extend(process_segment(after, base_style))
            return result

        # Handle **bold**
        bold_pattern = r'\*\*(.*?)\*\*'
        match = re.search(bold_pattern, s)
        if match:
            before = s[:match.start()]
            inner = match.group(1)
            after = s[match.end():]

            if before:
                result.extend(process_segment(before, base_style))
            bold_style = {**base_style, 'bold': True}
            result.extend(process_segment(inner, bold_style))
            if after:
                result.extend(process_segment(after, base_style))
            return result

        # Handle *italic* (but not **)
        italic_pattern = r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)'
        match = re.search(italic_pattern, s)
        if match:
            before = s[:match.start()]
            inner = match.group(1)
            after = s[match.end():]

            if before:
                result.extend(process_segment(before, base_style))
            italic_style = {**base_style, 'italic': True}
            result.extend(process_segment(inner, italic_style))
            if after:
                result.extend(process_segment(after, base_style))
            return result

        # Handle `code`
        code_pattern = r'`(.*?)`'
        match = re.search(code_pattern, s)
        if match:
            before = s[:match.start()]
            inner = match.group(1)
            after = s[match.end():]

            if before:
                result.extend(process_segment(before, base_style))
            code_style = {**base_style, 'code': True}
            result.append((inner, code_style))
            if after:
                result.extend(process_segment(after, base_style))
            return result

        # Clean up any remaining HTML tags
        clean = re.sub(r'<[^>]+>', '', s)
        clean = html.unescape(clean)
        if clean:
            result.append((clean, base_style))
        return result

    return process_segment(text)


def add_styled_runs(paragraph, text):
    """Add text with inline formatting to a paragraph."""
    segments = parse_inline(text)
    for seg_text, style in segments:
        run = paragraph.add_run(seg_text)
        if style.get('bold'):
            run.bold = True
        if style.get('italic'):
            run.italic = True
        if style.get('code'):
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        if 'color' in style:
            r, g, b = style['color']
            run.font.color.rgb = RGBColor(r, g, b)


def classify_line(line):
    """Classify a markdown line into its type and content."""
    stripped = line.strip()

    if not stripped:
        return ('blank', '', 0)

    # Headings
    heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2)
        return ('heading', text, level)

    # Horizontal rules
    if re.match(r'^---+\s*$', stripped):
        return ('hr', '', 0)

    # Blockquotes
    if stripped.startswith('>'):
        text = re.sub(r'^>\s*', '', stripped)
        return ('blockquote', text, 0)

    # Numbered list items
    num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
    if num_match:
        text = num_match.group(2)
        # Determine indent level from original line
        indent = len(line) - len(line.lstrip())
        indent_level = 0
        if indent > 0:
            indent_level = 1 if indent <= 4 else 2
        return ('numbered', text, indent_level)

    # Bullet list items
    bullet_match = re.match(r'^[-*]\s+(.*)', stripped)
    if bullet_match:
        text = bullet_match.group(1)
        indent = len(line) - len(line.lstrip())
        indent_level = 0
        if indent > 0:
            indent_level = 1 if indent <= 4 else 2
        return ('bullet', text, indent_level)

    # Regular text
    return ('text', stripped, 0)


def build_document(md_lines):
    """Build a DOCX from parsed markdown lines."""
    doc = Document()

    # Set up default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # Configure heading styles
    for i, (size, color_hex) in enumerate([
        (16, '1F3864'),  # Heading 1: Chapter
        (13, '2E75B6'),  # Heading 2: Section
        (11, '2E75B6'),  # Heading 3: Subsection
    ], start=1):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = RGBColor.from_string(color_hex)
        h_style.font.name = 'Calibri'
        h_style.paragraph_format.space_before = Pt(12 if i == 1 else 8)
        h_style.paragraph_format.space_after = Pt(4)

    # Set up auto-numbering
    num_id = setup_multilevel_numbering(doc)
    link_heading_style_to_numbering(doc, 'Heading 1', num_id, 0)
    link_heading_style_to_numbering(doc, 'Heading 2', num_id, 1)
    link_heading_style_to_numbering(doc, 'Heading 3', num_id, 2)

    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    run = title_para.add_run('FlexoGraph Dissertation Outline')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string('1F3864')

    # Add legend
    legend_para = doc.add_paragraph()
    legend_para.paragraph_format.space_after = Pt(2)
    run = legend_para.add_run('Legend: ')
    run.bold = True
    run.font.size = Pt(10)

    legend_items = [
        ('[MAIN]', 'Content on main branch (SIGMOD submission)'),
        ('[CANDIDACY]', 'Content only on candidacy branch'),
        ('[BOTH]', 'Content on both branches'),
        ('[TO WRITE]', 'Needs to be written or substantially expanded'),
        ('[DECIDE]', 'Decision needed'),
    ]
    for tag, desc in legend_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(tag)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f' — {desc}')
        run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Skip the title and legend lines in the source markdown
    # (we already added them above). Find the first ## heading.
    start = 0
    for idx, line in enumerate(md_lines):
        if line.strip().startswith('## '):
            start = idx
            break

    # Track state for chapters
    in_preliminary = False
    in_appendix = False

    i = start
    while i < len(md_lines):
        line = md_lines[i]
        line_type, content, level = classify_line(line)

        if line_type == 'blank' or line_type == 'hr':
            i += 1
            continue

        if line_type == 'heading':
            # Clean the heading text
            # Remove trailing tags like **[MAIN]** or [Estimate: ...]
            heading_text = content.strip()

            # Remove markdown bold from within heading
            heading_text = re.sub(r'\*\*\[', '[', heading_text)
            heading_text = re.sub(r'\]\*\*', ']', heading_text)

            # Remove any HTML tags
            heading_text = re.sub(r'<[^>]+>', '', heading_text)

            # Extract status tags for annotation
            status_tags = re.findall(r'\[(MAIN|CANDIDACY|BOTH|TO WRITE|DECIDE|TODO[^]]*|WRITTEN[^]]*|NEEDS[^]]*|EXPAND)\]', heading_text)

            # Extract estimate tags
            estimate_match = re.search(r'\[Estimate[^]]*\]', heading_text, re.IGNORECASE)
            estimate_text = estimate_match.group(0) if estimate_match else None

            # Remove tags and estimates from heading text
            heading_text = re.sub(r'\s*\[(?:MAIN|CANDIDACY|BOTH|TO WRITE|DECIDE|TODO[^]]*|WRITTEN[^]]*|NEEDS[^]]*|EXPAND)\]', '', heading_text)
            heading_text = re.sub(r'\s*\[Estimate[^]]*\]', '', heading_text, flags=re.IGNORECASE)
            heading_text = re.sub(r'\s*--\s*$', '', heading_text)  # trailing dashes

            # Map markdown heading levels to document structure
            # ## = Chapter (Heading 1), ### = Section (Heading 2), #### = Subsection (Heading 3)
            if level == 1:
                # Top-level title — skip (already handled)
                i += 1
                continue
            elif level == 2:
                doc_level = 1  # Heading 1
                # Check if this is "Preliminary Pages" — not a numbered chapter
                if 'Preliminary Pages' in heading_text:
                    p = doc.add_paragraph()
                    run = p.add_run(heading_text.strip())
                    run.bold = True
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor.from_string('1F3864')
                    p.paragraph_format.space_before = Pt(12)
                    in_preliminary = True
                    i += 1
                    continue
                elif 'Appendix' in heading_text:
                    p = doc.add_paragraph()
                    run = p.add_run(heading_text.strip())
                    run.bold = True
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor.from_string('1F3864')
                    p.paragraph_format.space_before = Pt(12)
                    in_appendix = True
                    i += 1
                    continue
                else:
                    in_preliminary = False
                    in_appendix = False
                # Strip "Chapter N:" prefix — numbering is auto
                heading_text = strip_chapter_prefix(heading_text)
            elif level == 3:
                doc_level = 2  # Heading 2
                heading_text = strip_section_number(heading_text)
            elif level == 4:
                doc_level = 3  # Heading 3
                heading_text = strip_section_number(heading_text)
            else:
                doc_level = 3

            heading_text = heading_text.strip()
            if not heading_text:
                i += 1
                continue

            # Add the heading
            if in_preliminary:
                # Preliminary pages items as bullets
                p = doc.add_paragraph(style='List Bullet')
                add_styled_runs(p, heading_text)
            elif in_appendix:
                # Appendix subsections — styled but not auto-numbered
                p = doc.add_paragraph()
                run = p.add_run(heading_text.strip())
                run.bold = True
                sizes = {2: Pt(13), 3: Pt(11)}
                run.font.size = sizes.get(doc_level, Pt(11))
                run.font.color.rgb = RGBColor.from_string('2E75B6')
                p.paragraph_format.space_before = Pt(8)
            else:
                p = doc.add_heading(heading_text, level=doc_level)

            # Add status tags as colored annotation after heading
            if status_tags:
                tag_str = ' [' + ', '.join(status_tags) + ']'
                run = p.add_run(tag_str)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                run.bold = False

            if estimate_text:
                run = p.add_run(f'  {estimate_text}')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                run.italic = True
                run.bold = False

            i += 1
            continue

        if line_type == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            # Style blockquotes with a left border color indication
            add_styled_runs(p, content)
            # Make the text slightly smaller and italic for notes/decisions
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(10)
                if not run.font.color.rgb:
                    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            i += 1
            continue

        if line_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            if level > 0:
                p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
                p.style = doc.styles['List Bullet 2'] if level == 1 else doc.styles['List Bullet 3']
            add_styled_runs(p, content)
            i += 1
            continue

        if line_type == 'numbered':
            p = doc.add_paragraph(style='List Number')
            if level > 0:
                p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
                try:
                    p.style = doc.styles['List Number 2'] if level == 1 else doc.styles['List Number 3']
                except KeyError:
                    pass
            add_styled_runs(p, content)
            i += 1
            continue

        if line_type == 'text':
            p = doc.add_paragraph()
            add_styled_runs(p, content)
            i += 1
            continue

        i += 1

    return doc


def main():
    with open('/Users/puneet/Documents/thesis/thesis-outline.md', 'r') as f:
        lines = f.readlines()

    # Remove trailing newlines but preserve indentation
    lines = [line.rstrip('\n') for line in lines]

    doc = build_document(lines)

    output_path = '/Users/puneet/Documents/thesis/thesis-outline.docx'
    doc.save(output_path)
    print(f'Saved to {output_path}')


if __name__ == '__main__':
    main()
